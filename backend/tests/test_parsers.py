"""Tests for document parsers (PDF, DOCX, CSV, TXT)."""
from __future__ import annotations

import io
import struct

import pytest

from app.services.parsers import parse_document


# ─── TXT parsing ─────────────────────────────────────────────────────────────

class TestTxtParser:
    def test_plain_text_roundtrip(self):
        content = "Hello, world!\nSecond line."
        result = parse_document(content.encode("utf-8"), "notes.txt")
        assert result == content

    def test_utf8_with_special_chars(self):
        content = "Héllo café résumé"
        result = parse_document(content.encode("utf-8"), "special.txt")
        assert "café" in result

    def test_empty_text(self):
        result = parse_document(b"", "empty.txt")
        assert result == ""

    def test_binary_fallback(self):
        """Malformed UTF-8 bytes should be replaced, not crash."""
        bad_bytes = b"Hello \xff\xfe world"
        result = parse_document(bad_bytes, "bad.txt")
        assert "Hello" in result
        assert "world" in result


# ─── CSV parsing ─────────────────────────────────────────────────────────────

class TestCsvParser:
    def test_simple_csv(self):
        csv_data = "name,age,city\nAlice,30,NYC\nBob,25,LA\n"
        result = parse_document(csv_data.encode(), "people.csv")
        assert "Alice" in result
        assert "name: Alice" in result
        assert "age: 30" in result

    def test_csv_with_missing_values(self):
        """Rows with NaN values should still include the non-NaN columns."""
        csv_data = "name,score,grade\nAlice,95,A\nBob,,B\nCharlie,87,\n"
        result = parse_document(csv_data.encode(), "grades.csv")
        lines = result.strip().split("\n")
        assert len(lines) == 3
        # Bob's row should have name and grade but no score
        bob_line = lines[1]
        assert "name: Bob" in bob_line
        assert "grade: B" in bob_line
        # Charlie's row should have name and score but no grade
        charlie_line = lines[2]
        assert "name: Charlie" in charlie_line
        assert "score: 87" in charlie_line

    def test_csv_preserves_all_columns(self):
        """All non-NaN column headers should appear in output."""
        csv_data = "id,product,price\n1,Widget,9.99\n2,Gadget,19.99\n"
        result = parse_document(csv_data.encode(), "products.csv")
        assert "id: 1" in result
        assert "product: Widget" in result
        assert "price: 9.99" in result

    def test_single_column_csv(self):
        csv_data = "email\nalice@test.com\nbob@test.com\n"
        result = parse_document(csv_data.encode(), "emails.csv")
        assert "email: alice@test.com" in result


# ─── DOCX parsing ────────────────────────────────────────────────────────────

class TestDocxParser:
    def test_docx_basic(self):
        """Create a minimal DOCX in memory and parse it."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph with details.")
        doc.add_paragraph("")  # blank paragraph — should be filtered
        doc.add_paragraph("Third paragraph.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = parse_document(buf.read(), "test.docx")
        assert "First paragraph." in result
        assert "Second paragraph with details." in result
        assert "Third paragraph." in result


# ─── PDF parsing ─────────────────────────────────────────────────────────────

class TestPdfParser:
    def test_pdf_basic(self):
        """Create a minimal PDF with PyMuPDF and parse it."""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from page one.")
        page2 = doc.new_page()
        page2.insert_text((72, 72), "Content on page two.")

        pdf_bytes = doc.tobytes()
        doc.close()

        result = parse_document(pdf_bytes, "test.pdf")
        assert "[Page 1]" in result
        assert "Hello from page one." in result
        assert "[Page 2]" in result
        assert "Content on page two." in result

    def test_pdf_blank_page_warning(self, caplog):
        """PDFs with blank pages should log a warning but still parse."""
        import fitz

        doc = fitz.open()
        doc.new_page()  # blank page
        page2 = doc.new_page()
        page2.insert_text((72, 72), "Only content page.")

        pdf_bytes = doc.tobytes()
        doc.close()

        import logging
        with caplog.at_level(logging.WARNING):
            result = parse_document(pdf_bytes, "partial.pdf")

        assert "Only content page." in result
        assert "blank pages" in caplog.text

    def test_pdf_all_blank_raises(self):
        """A PDF with no extractable text should raise ValueError."""
        import fitz

        doc = fitz.open()
        doc.new_page()  # blank
        doc.new_page()  # blank

        pdf_bytes = doc.tobytes()
        doc.close()

        with pytest.raises(ValueError, match="No text could be extracted"):
            parse_document(pdf_bytes, "scanned.pdf")


# ─── Unsupported type ────────────────────────────────────────────────────────

class TestUnsupportedType:
    def test_unsupported_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_document(b"data", "image.png")

    def test_unknown_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_document(b"data", "file.xyz")
